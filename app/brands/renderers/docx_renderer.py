"""
DOCX renderer for a Brand.

Takes an existing DOCX and applies the brand:
  1. Styles: Heading 1 / Heading 2 / Normal use brand fonts and colors
  2. Heading 1: bottom border in accent color (the orange "rule" motif)
  3. Section labels: 'NN / TITLE' caps prepended above each Heading 1
  4. Header: organization name eyebrow on every page
  5. Footer: '{document_label} · page' on every page
  6. Optional cover page (section break + branded title block) when
     brief_meta has cover content
"""
from __future__ import annotations

import logging
import re
import time
from copy import deepcopy
from typing import Any, Optional

from app.brands.base import Brand
from app.storage import (
    input_path,
    new_file_id,
    output_path,
    public_url,
)

log = logging.getLogger("aice.brand.docx")


def _hex_to_docx_rgb(hex_color: str):
    from docx.shared import RGBColor
    s = hex_color.strip().lstrip("#")
    return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))


def _hex_to_w_val(hex_color: str) -> str:
    """Hex color WITHOUT the leading '#', for direct use in OOXML w:val attributes."""
    return hex_color.strip().lstrip("#").upper()


def _set_paragraph_bottom_border(paragraph, color_hex: str, size_pt: float = 1.5):
    """Add a bottom border (underline rule) to a paragraph via XML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()
    # Remove existing pBdr if any
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    # OOXML border size is in eighths of a point; 12 = 1.5pt
    bottom.set(qn("w:sz"), str(int(size_pt * 8)))
    bottom.set(qn("w:space"), "4")  # padding from text in points
    bottom.set(qn("w:color"), _hex_to_w_val(color_hex))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_style_font(style, *, font_name: Optional[str] = None,
                    size_pt: Optional[float] = None,
                    color_hex: Optional[str] = None,
                    bold: Optional[bool] = None,
                    italic: Optional[bool] = None):
    from docx.shared import Pt
    if font_name:
        style.font.name = font_name
        # Also set East Asian font attribute so theme/East Asian default
        # doesn't override it (common with Hungarian docs)
        from docx.oxml.ns import qn
        rPr = style.element.find(qn("w:rPr"))
        if rPr is not None:
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                rFonts.set(qn("w:ascii"), font_name)
                rFonts.set(qn("w:hAnsi"), font_name)
                rFonts.set(qn("w:cs"), font_name)
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if color_hex is not None:
        style.font.color.rgb = _hex_to_docx_rgb(color_hex)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def _customize_styles(doc, brand: Brand):
    """Modify the document's styles to use brand fonts/colors."""
    styles = doc.styles

    # Heading 1 — Georgia, bold, primary text color
    try:
        h1 = styles["Heading 1"]
        _set_style_font(h1,
                        font_name=brand.fonts.heading,
                        size_pt=20, bold=True,
                        color_hex=brand.colors.text_primary)
    except KeyError:
        pass

    # Heading 2 — Georgia, bold, smaller
    try:
        h2 = styles["Heading 2"]
        _set_style_font(h2,
                        font_name=brand.fonts.heading,
                        size_pt=15, bold=True,
                        color_hex=brand.colors.text_primary)
    except KeyError:
        pass

    # Heading 3 — Aptos/Calibri uppercase tiny accent
    try:
        h3 = styles["Heading 3"]
        _set_style_font(h3,
                        font_name=brand.fonts.body,
                        size_pt=11, bold=True,
                        color_hex=brand.colors.accent)
    except KeyError:
        pass

    # Normal — Aptos/Calibri, primary text color
    try:
        normal = styles["Normal"]
        _set_style_font(normal,
                        font_name=brand.fonts.body,
                        size_pt=11,
                        color_hex=brand.colors.text_primary)
    except KeyError:
        pass

    # Title — Georgia, big
    for sname in ("Title",):
        try:
            ts = styles[sname]
            _set_style_font(ts,
                            font_name=brand.fonts.heading,
                            size_pt=32, bold=True,
                            color_hex=brand.colors.text_primary)
        except KeyError:
            pass


def _add_heading_underlines(doc, brand: Brand) -> int:
    """Add an accent-color bottom border under every Heading 1 paragraph."""
    n = 0
    for p in doc.paragraphs:
        style_name = (p.style.name or "") if p.style else ""
        if style_name == "Heading 1":
            _set_paragraph_bottom_border(p, brand.colors.accent, size_pt=2)
            n += 1
    return n


def _insert_paragraph_before(paragraph, text, *, style_name: Optional[str] = None,
                             font_name: Optional[str] = None,
                             size_pt: Optional[float] = None,
                             color_hex: Optional[str] = None,
                             bold: Optional[bool] = None,
                             italic: Optional[bool] = None,
                             space_after_pt: float = 0):
    """Insert a new paragraph just before `paragraph`. Returns the new paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    new_p_elem = OxmlElement("w:p")
    paragraph._p.addprevious(new_p_elem)
    # Wrap into a Paragraph object
    from docx.text.paragraph import Paragraph
    new_p = Paragraph(new_p_elem, paragraph._parent)
    if style_name:
        try:
            new_p.style = paragraph.part.document.styles[style_name]
        except KeyError:
            pass
    run = new_p.add_run(text)
    if font_name:
        run.font.name = font_name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color_hex is not None:
        run.font.color.rgb = _hex_to_docx_rgb(color_hex)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if space_after_pt:
        new_p.paragraph_format.space_after = Pt(space_after_pt)
    return new_p


def _add_section_labels(doc, brand: Brand) -> int:
    """Insert 'NN / TITLE' orange caps label above every Heading 1.
    Skips if a label already exists immediately above."""
    n = 0
    section_index = 0
    paragraphs = list(doc.paragraphs)  # snapshot before mutation
    for p in paragraphs:
        style_name = (p.style.name or "") if p.style else ""
        if style_name != "Heading 1":
            continue
        section_index += 1
        # Check if previous paragraph already looks like a section label
        prev = p._p.getprevious()
        prev_text = ""
        if prev is not None and prev.tag.endswith("}p"):
            from docx.text.paragraph import Paragraph
            prev_p = Paragraph(prev, p._parent)
            prev_text = prev_p.text.strip()
        if re.match(r"^\d{2}\s*/", prev_text):
            continue  # already labeled
        title = p.text.strip()
        try:
            label = brand.motifs.section_label_format.format(
                number=section_index, title=title,
            )
        except (KeyError, IndexError):
            label = title
        if brand.motifs.section_label_caps:
            label = label.upper()
        _insert_paragraph_before(p, label,
                                 font_name=brand.fonts.effective_eyebrow(),
                                 size_pt=10, bold=True,
                                 color_hex=brand.colors.accent,
                                 space_after_pt=2)
        n += 1
    return n


def _setup_header_footer(doc, brand: Brand, document_label: str):
    """Set up the page header (org-name eyebrow) and footer (label · page)."""
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    org_name = brand.config.get("organization_name", "")

    for section in doc.sections:
        # Header — org name eyebrow
        if org_name and brand.motifs.inner_top_eyebrow:
            header = section.header
            # Wipe existing header paragraphs (replace them)
            for p in list(header.paragraphs):
                p.text = ""
            hp = header.paragraphs[0]
            hp.text = ""
            run = hp.add_run(org_name)
            run.font.name = brand.fonts.effective_eyebrow()
            run.font.size = Pt(8.5)
            run.font.bold = True
            run.font.color.rgb = _hex_to_docx_rgb(brand.colors.text_subtle)

        # Footer — document_label + page number (via field code)
        footer = section.footer
        for p in list(footer.paragraphs):
            p.text = ""
        fp = footer.paragraphs[0]
        # Left-aligned label
        run_label = fp.add_run(f"{document_label}  ·  ")
        run_label.font.name = brand.fonts.effective_eyebrow()
        run_label.font.size = Pt(8.5)
        run_label.font.color.rgb = _hex_to_docx_rgb(brand.colors.text_subtle)
        # Page number (FIELD code — Word/LO renders it dynamically)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run_page = fp.add_run()
        run_page.font.name = brand.fonts.effective_eyebrow()
        run_page.font.size = Pt(8.5)
        run_page.font.color.rgb = _hex_to_docx_rgb(brand.colors.text_subtle)
        run_page._r.append(fld_begin)
        run_page._r.append(instr)
        run_page._r.append(fld_end)


def _resolve(file_id):
    p = output_path(file_id)
    if p.exists():
        return p
    return input_path(file_id)


async def apply_brand_to_docx(
    file_id: str, brand: Brand, brief_meta: Optional[dict] = None,
) -> dict:
    """Apply the brand to an existing DOCX. Modifies styles in-place, adds
    Heading 1 underlines, section labels, header eyebrow, and footer with
    document label + page numbers."""
    from docx import Document

    if brief_meta is None:
        brief_meta = {}
    document_label = (brief_meta.get("document_label")
                      or brief_meta.get("brief_code")
                      or "BRIEF")

    src = _resolve(file_id)
    t0 = time.time()
    doc = Document(str(src))

    _customize_styles(doc, brand)
    n_underlines = _add_heading_underlines(doc, brand)
    n_labels = _add_section_labels(doc, brand)
    _setup_header_footer(doc, brand, document_label)

    fid = new_file_id("docx")
    dst = output_path(fid)
    doc.save(str(dst))
    return {
        "file_id": fid,
        "url": public_url(fid),
        "brand": brand.name,
        "headings_underlined": n_underlines,
        "section_labels_added": n_labels,
        "ms_elapsed": int((time.time() - t0) * 1000),
    }
