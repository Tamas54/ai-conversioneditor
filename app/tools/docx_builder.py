"""
DOCX builder — agentic document construction tools (analogue of pptx_builder).

The agent uses these to BUILD a new DOCX from scratch, slide-by-slide style:
  1. docx_create(title?) — start an empty document
  2. docx_add_heading / add_paragraph / add_form_field / add_numbered_list /
     add_bulleted_list / add_signature_block / add_table / add_page_break
  3. docx_set_page_setup — orientation + margins
  4. (Then optionally pass file_id to brand_apply for branded chrome)

Each call produces a new file_id; the agent threads them via state.current_fid.
"""
from __future__ import annotations

import logging
import time
from copy import deepcopy
from pathlib import Path
from typing import Any, Optional

from app.storage import (
    input_path,
    new_file_id,
    output_path,
    public_url,
)

log = logging.getLogger("aice.docx_builder")


def _resolve(file_id: str) -> Path:
    p = output_path(file_id)
    if p.exists():
        return p
    p = input_path(file_id)
    if p.exists():
        return p
    raise FileNotFoundError(f"file_id not found: {file_id}")


def _new_docx_output(
    *, source_file_id: str | None = None, operation: str | None = None,
    label: str | None = None,
) -> tuple[str, Path]:
    fid = new_file_id("docx")
    if source_file_id or operation or label:
        from app.storage import write_meta
        write_meta(
            fid, source_file_id=source_file_id,
            operation=operation, label=label,
        )
    return fid, output_path(fid)


def _hex_to_rgb(hex_color: str):
    from docx.shared import RGBColor
    s = hex_color.strip().lstrip("#")
    return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))


# ============================================================
# Builder primitives
# ============================================================


async def docx_create(title: Optional[str] = None) -> dict[str, Any]:
    """Create a new empty DOCX. If `title` provided, adds it as a centered
    bold Title-styled paragraph at the top."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    t0 = time.time()
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(16)
        run.font.bold = True

    fid, dst = _new_docx_output()
    doc.save(str(dst))
    return {
        "file_id": fid,
        "url": public_url(fid),
        "ms_elapsed": int((time.time() - t0) * 1000),
    }


async def docx_set_page_setup(
    file_id: str,
    *,
    orientation: str = "portrait",   # 'portrait' | 'landscape'
    page_size: str = "A4",            # 'A4' | 'Letter' | 'Legal'
    margin_cm: float = 2.5,
) -> dict[str, Any]:
    from docx import Document
    from docx.shared import Cm
    from docx.enum.section import WD_ORIENTATION

    src = _resolve(file_id)
    t0 = time.time()
    doc = Document(str(src))
    section = doc.sections[0]
    sizes = {"A4": (21.0, 29.7), "Letter": (21.59, 27.94), "Legal": (21.59, 35.56)}
    w_cm, h_cm = sizes.get(page_size, sizes["A4"])
    if orientation == "landscape":
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = Cm(h_cm)
        section.page_height = Cm(w_cm)
    else:
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width = Cm(w_cm)
        section.page_height = Cm(h_cm)
    section.top_margin = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)
    section.left_margin = Cm(margin_cm)
    section.right_margin = Cm(margin_cm)

    fid, dst = _new_docx_output()
    doc.save(str(dst))
    return {"file_id": fid, "url": public_url(fid), "ms_elapsed": int((time.time() - t0) * 1000)}


def _resolve_alignment(name: Optional[str]):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get((name or "").lower())


async def docx_add_heading(
    file_id: str, text: str, *, level: int = 1,
    alignment: Optional[str] = None,
) -> dict[str, Any]:
    """Add a heading. level ∈ {1,2,3}. Uses Word's built-in 'Heading N' style
    so brand_apply later restyles it consistently."""
    from docx import Document

    src = _resolve(file_id)
    t0 = time.time()
    doc = Document(str(src))
    h = doc.add_heading(text, level=max(1, min(level, 3)))
    align = _resolve_alignment(alignment)
    if align is not None:
        h.alignment = align
    fid, dst = _new_docx_output()
    doc.save(str(dst))
    return {"file_id": fid, "url": public_url(fid), "ms_elapsed": int((time.time() - t0) * 1000)}


async def docx_add_paragraph(
    file_id: str, text: str, *,
    bold: bool = False,
    italic: bool = False,
    alignment: Optional[str] = None,    # 'left'|'center'|'right'|'justify'
    font_size_pt: Optional[float] = None,
    font_family: Optional[str] = None,
    color: Optional[str] = None,        # hex like '#0E1726'
    space_after_pt: Optional[float] = None,
    indent_cm: Optional[float] = None,
) -> dict[str, Any]:
    """Add a body paragraph with optional formatting overrides."""
    from docx import Document
    from docx.shared import Pt, Cm

    src = _resolve(file_id)
    t0 = time.time()
    doc = Document(str(src))
    p = doc.add_paragraph()
    align = _resolve_alignment(alignment)
    if align is not None:
        p.alignment = align
    if indent_cm is not None:
        p.paragraph_format.left_indent = Cm(indent_cm)
    if space_after_pt is not None:
        p.paragraph_format.space_after = Pt(space_after_pt)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)
    if font_family:
        run.font.name = font_family
    if color:
        run.font.color.rgb = _hex_to_rgb(color)
    fid, dst = _new_docx_output()
    doc.save(str(dst))
    return {"file_id": fid, "url": public_url(fid), "ms_elapsed": int((time.time() - t0) * 1000)}


async def docx_add_form_field(
    file_id: str, label: str, *,
    dots: int = 80, label_bold: bool = True,
) -> dict[str, Any]:
    """Add a form field line: 'Label: ........................................'."""
    from docx import Document
    from docx.shared import Pt

    src = _resolve(file_id)
    t0 = time.time()
    doc = Document(str(src))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(label + " ")
    run.bold = label_bold
    run2 = p.add_run("." * dots)
    fid, dst = _new_docx_output()
    doc.save(str(dst))
    return {"file_id": fid, "url": public_url(fid), "ms_elapsed": int((time.time() - t0) * 1000)}


async def docx_add_numbered_list(
    file_id: str, items: list[str], *,
    indent_cm: float = 0.8,
) -> dict[str, Any]:
    """Add a numbered list. Uses '1.   item' style with indented paragraphs."""
    from docx import Document
    from docx.shared import Pt, Cm

    src = _resolve(file_id)
    t0 = time.time()
    doc = Document(str(src))
    for i, it in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(f"{i}.   {it}")
    fid, dst = _new_docx_output()
    doc.save(str(dst))
    return {
        "file_id": fid, "url": public_url(fid),
        "items_added": len(items),
        "ms_elapsed": int((time.time() - t0) * 1000),
    }


async def docx_add_bulleted_list(
    file_id: str, items: list[str], *,
    indent_cm: float = 0.8,
) -> dict[str, Any]:
    from docx import Document
    from docx.shared import Pt, Cm

    src = _resolve(file_id)
    t0 = time.time()
    doc = Document(str(src))
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.paragraph_format.space_after = Pt(4)
        # If the style isn't available, prefix with "• "
        if not p.runs:
            p.add_run(f"• {it}")
        else:
            p.add_run(it)
    fid, dst = _new_docx_output()
    doc.save(str(dst))
    return {
        "file_id": fid, "url": public_url(fid),
        "items_added": len(items),
        "ms_elapsed": int((time.time() - t0) * 1000),
    }


async def docx_add_signature_block(
    file_id: str,
    label: str = "aláírás",
    *,
    alignment: str = "right",
    dots: int = 50,
    space_before_pt: float = 24,
) -> dict[str, Any]:
    """Add a dotted signature line + small italic caption underneath."""
    from docx import Document
    from docx.shared import Pt

    src = _resolve(file_id)
    t0 = time.time()
    doc = Document(str(src))
    align = _resolve_alignment(alignment) or _resolve_alignment("right")

    # Optional spacer paragraph
    if space_before_pt:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(space_before_pt)

    sig = doc.add_paragraph()
    sig.alignment = align
    sig.paragraph_format.space_after = Pt(2)
    sig.add_run("." * dots)

    cap = doc.add_paragraph()
    cap.alignment = align
    cr = cap.add_run(label)
    cr.italic = True
    cr.font.size = Pt(10)

    fid, dst = _new_docx_output()
    doc.save(str(dst))
    return {"file_id": fid, "url": public_url(fid), "ms_elapsed": int((time.time() - t0) * 1000)}


async def docx_add_table(
    file_id: str,
    headers: list[str],
    rows: list[list[str]],
    *,
    header_color: str = "#1F3A5F",
    header_text_color: str = "#FFFFFF",
    row_band_color: Optional[str] = "#F4E7C5",
) -> dict[str, Any]:
    """Add a styled table with bold colored header and optional row banding."""
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    src = _resolve(file_id)
    t0 = time.time()
    doc = Document(str(src))

    n_cols = max(1, len(headers))
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"

    def _cell_shade(cell, hex_color: str):
        s = hex_color.lstrip("#")
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        # Remove existing shd if any
        for old in tcPr.findall(qn("w:shd")):
            tcPr.remove(old)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), s.upper())
        tcPr.append(shd)

    # Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers[:n_cols]):
        cell = hdr_cells[i]
        cell.text = ""
        _cell_shade(cell, header_color)
        p = cell.paragraphs[0]
        run = p.add_run(str(h))
        run.bold = True
        run.font.color.rgb = _hex_to_rgb(header_text_color)
        run.font.size = Pt(11)

    # Data rows
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        if row_band_color and r_idx % 2 == 0:
            for c in cells:
                _cell_shade(c, row_band_color)
        for c_idx in range(n_cols):
            cell = cells[c_idx]
            val = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val) if val is not None else "")
            run.font.size = Pt(10)

    fid, dst = _new_docx_output()
    doc.save(str(dst))
    return {
        "file_id": fid, "url": public_url(fid),
        "rows": len(rows), "cols": n_cols,
        "ms_elapsed": int((time.time() - t0) * 1000),
    }


async def docx_add_image_inline(
    file_id: str, image_file_id: str, *,
    width_inches: Optional[float] = None,
    alignment: str = "center",
    caption: Optional[str] = None,
) -> dict[str, Any]:
    """Add an image inline with optional caption below it."""
    from docx import Document
    from docx.shared import Inches, Pt

    src = _resolve(file_id)
    img = _resolve(image_file_id)
    t0 = time.time()
    doc = Document(str(src))

    p = doc.add_paragraph()
    p.alignment = _resolve_alignment(alignment) or _resolve_alignment("center")
    run = p.add_run()
    kwargs = {"width": Inches(width_inches)} if width_inches else {}
    run.add_picture(str(img), **kwargs)

    if caption:
        cp = doc.add_paragraph()
        cp.alignment = _resolve_alignment(alignment) or _resolve_alignment("center")
        cr = cp.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(10)

    fid, dst = _new_docx_output()
    doc.save(str(dst))
    return {"file_id": fid, "url": public_url(fid), "ms_elapsed": int((time.time() - t0) * 1000)}


async def docx_add_page_break(file_id: str) -> dict[str, Any]:
    from docx import Document
    src = _resolve(file_id)
    t0 = time.time()
    doc = Document(str(src))
    doc.add_page_break()
    fid, dst = _new_docx_output()
    doc.save(str(dst))
    return {"file_id": fid, "url": public_url(fid), "ms_elapsed": int((time.time() - t0) * 1000)}
