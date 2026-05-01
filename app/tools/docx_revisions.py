"""
DOCX revisions — tracked changes + comments.

Implements the OOXML revision-tracking surface that nobody in OSS does well:
  * <w:ins author="..." date="..." id="N">  — tracked insertion
  * <w:del author="..." date="..." id="N">  — tracked deletion (uses <w:delText>)
  * Word's "review pane" lights up automatically when these elements exist
  * Comments via word/comments.xml + commentRangeStart/End anchors

Use cases:
  * Legal/academic doc review: agent proposes edits as redlines, human approves
  * Two-author collaboration: agent's changes attributed properly
  * Comment-driven workflow: agent reads comments, proposes per-comment edits

Tools:
  - docx_track_replace_paragraph: replace a paragraph's text as <w:del>+<w:ins>
  - docx_track_insert_after: insert a new paragraph after anchor as <w:ins>
  - docx_track_delete_paragraph: mark a paragraph as <w:del>
  - docx_accept_all_revisions: collapse <w:ins>→keep, <w:del>→remove
  - docx_reject_all_revisions: collapse <w:ins>→remove, <w:del>→keep
  - docx_list_revisions: enumerate all <w:ins>/<w:del> with author + date + text
  - docx_add_comment: anchor a comment to a paragraph
  - docx_list_comments: enumerate all comments with id + author + text + anchor
  - docx_reply_to_comment: thread a reply via commentsExtended.xml
  - docx_resolve_comment: mark a comment thread as resolved
"""
from __future__ import annotations

import logging
import re
import time
import zipfile
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

from app.storage import (
    input_path,
    new_file_id,
    output_path,
    public_url,
)

log = logging.getLogger("aice.docx_revisions")

# OOXML namespaces
NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS_CT = "http://schemas.openxmlformats.org/package/2006/content-types"
NS_PR = "http://schemas.openxmlformats.org/package/2006/relationships"


def _resolve(file_id: str) -> Path:
    p = output_path(file_id)
    if p.exists():
        return p
    return input_path(file_id)


def _new_docx_output(
    *, source_file_id: str | None = None, operation: str | None = None,
    label: str | None = None,
) -> tuple[str, Path]:
    fid = new_file_id("docx")
    if source_file_id or operation or label:
        from app.storage import write_meta
        write_meta(
            fid, source_file_id=source_file_id,
            operation=operation, label=label,
        )
    return fid, output_path(fid)


def _now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def qn(tag: str) -> str:
    """Resolve a w:tag to a fully namespaced ElementTree tag."""
    if ":" in tag:
        prefix, local = tag.split(":", 1)
        ns = {"w": NS_W, "r": NS_R}.get(prefix, "")
        return f"{{{ns}}}{local}"
    return f"{{{NS_W}}}{tag}"


def _next_revision_id(doc) -> int:
    """Find the highest existing w:id in the document, return next."""
    body = doc.element.body
    max_id = 0
    for el in body.iter():
        wid = el.get(qn("w:id"))
        if wid:
            try:
                max_id = max(max_id, int(wid))
            except ValueError:
                pass
    return max_id + 1


def _make_run(text: str, *, del_text: bool = False) -> Any:
    """Build a <w:r><w:t>text</w:t></w:r> (or <w:delText>) element via lxml."""
    from lxml import etree
    r = etree.Element(qn("w:r"))
    t_tag = qn("w:delText") if del_text else qn("w:t")
    t = etree.SubElement(r, t_tag)
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


def _wrap_in_ins(runs: list, *, author: str, date: str, rid: int) -> Any:
    from lxml import etree
    ins = etree.Element(qn("w:ins"))
    ins.set(qn("w:id"), str(rid))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date)
    for r in runs:
        ins.append(r)
    return ins


def _wrap_in_del(runs: list, *, author: str, date: str, rid: int) -> Any:
    from lxml import etree
    d = etree.Element(qn("w:del"))
    d.set(qn("w:id"), str(rid))
    d.set(qn("w:author"), author)
    d.set(qn("w:date"), date)
    for r in runs:
        d.append(r)
    return d


# ============================================================
# Tracked-change edit operations
# ============================================================


async def docx_track_replace_paragraph(
    file_id: str, anchor: str, new_text: str,
    *, author: str = "Kimi K2.6", change_id: Optional[int] = None,
) -> dict[str, Any]:
    """Replace the text of the paragraph whose body equals `anchor`. Old runs
    become <w:del>; new_text becomes <w:ins>. The original paragraph keeps its
    style. Word's review pane will show this as a tracked replacement."""
    from docx import Document
    from lxml import etree

    src = _resolve(file_id)
    t0 = time.time()
    doc = Document(str(src))
    target = anchor.strip()
    rid = change_id or _next_revision_id(doc)
    date = _now_iso()
    n = 0

    for p in doc.paragraphs:
        if p.text.strip() != target:
            continue
        p_elem = p._element

        # Collect existing runs (skip rPr and pPr — they're paragraph-level)
        existing_runs = []
        for child in list(p_elem):
            if child.tag == qn("w:r"):
                existing_runs.append(child)

        # Convert existing runs' <w:t> → <w:delText>, wrap them in <w:del>
        del_runs = []
        for r in existing_runs:
            r_clone = deepcopy(r)
            for t in r_clone.findall(qn("w:t")):
                # rename to delText
                t.tag = qn("w:delText")
            del_runs.append(r_clone)
            p_elem.remove(r)

        del_wrapper = _wrap_in_del(del_runs, author=author, date=date, rid=rid)
        p_elem.append(del_wrapper)

        # Insert new text in <w:ins>
        ins_wrapper = _wrap_in_ins(
            [_make_run(new_text)], author=author, date=date, rid=rid + 1,
        )
        p_elem.append(ins_wrapper)

        n += 1
        rid += 2

    if n == 0:
        raise ValueError(f"no paragraph found with text: {anchor!r}")

    fid, dst = _new_docx_output()
    doc.save(str(dst))
    return {
        "file_id": fid, "url": public_url(fid),
        "paragraphs_revised": n, "author": author,
        "ms_elapsed": int((time.time() - t0) * 1000),
    }


async def docx_track_insert_after(
    file_id: str, anchor: str, new_text: str,
    *, author: str = "Kimi K2.6",
) -> dict[str, Any]:
    """Insert `new_text` as a new paragraph immediately after `anchor`,
    marked as <w:ins>."""
    from docx import Document
    from lxml import etree

    src = _resolve(file_id)
    t0 = time.time()
    doc = Document(str(src))
    target = anchor.strip()
    rid = _next_revision_id(doc)
    date = _now_iso()

    for p in doc.paragraphs:
        if p.text.strip() != target:
            continue
        p_elem = p._element

        # Build a new paragraph element with the same pPr (style)
        new_p = etree.Element(qn("w:p"))
        # Copy paragraph properties for style continuity
        pPr = p_elem.find(qn("w:pPr"))
        if pPr is not None:
            new_p.append(deepcopy(pPr))
        ins_wrapper = _wrap_in_ins(
            [_make_run(new_text)], author=author, date=date, rid=rid,
        )
        new_p.append(ins_wrapper)

        p_elem.addnext(new_p)

        fid, dst = _new_docx_output()
        doc.save(str(dst))
        return {
            "file_id": fid, "url": public_url(fid),
            "author": author, "revision_id": rid,
            "ms_elapsed": int((time.time() - t0) * 1000),
        }
    raise ValueError(f"anchor paragraph not found: {anchor!r}")


async def docx_track_delete_paragraph(
    file_id: str, anchor: str, *, author: str = "Kimi K2.6",
) -> dict[str, Any]:
    """Mark the paragraph whose text equals `anchor` for deletion (tracked).
    Word will show it strikethrough in the review pane."""
    from docx import Document
    from lxml import etree

    src = _resolve(file_id)
    t0 = time.time()
    doc = Document(str(src))
    target = anchor.strip()
    rid = _next_revision_id(doc)
    date = _now_iso()
    n = 0

    for p in doc.paragraphs:
        if p.text.strip() != target:
            continue
        p_elem = p._element
        runs_to_del = []
        for child in list(p_elem):
            if child.tag == qn("w:r"):
                clone = deepcopy(child)
                for t in clone.findall(qn("w:t")):
                    t.tag = qn("w:delText")
                runs_to_del.append(clone)
                p_elem.remove(child)
        del_wrapper = _wrap_in_del(runs_to_del, author=author, date=date, rid=rid)
        p_elem.append(del_wrapper)

        # Mark paragraph mark itself as deleted via <w:rPr><w:del .../></w:rPr> on pPr's rPr
        pPr = p_elem.find(qn("w:pPr"))
        if pPr is None:
            pPr = etree.SubElement(p_elem, qn("w:pPr"))
            p_elem.insert(0, pPr)
            p_elem.remove(pPr)
            p_elem.insert(0, pPr)
        rPr = pPr.find(qn("w:rPr"))
        if rPr is None:
            rPr = etree.SubElement(pPr, qn("w:rPr"))
        del_mark = etree.SubElement(rPr, qn("w:del"))
        del_mark.set(qn("w:id"), str(rid + 1))
        del_mark.set(qn("w:author"), author)
        del_mark.set(qn("w:date"), date)
        n += 1
        rid += 2

    if n == 0:
        raise ValueError(f"no paragraph found with text: {anchor!r}")

    fid, dst = _new_docx_output()
    doc.save(str(dst))
    return {
        "file_id": fid, "url": public_url(fid),
        "paragraphs_marked": n, "author": author,
        "ms_elapsed": int((time.time() - t0) * 1000),
    }


# ============================================================
# Accept / reject operations
# ============================================================


async def docx_accept_all_revisions(file_id: str) -> dict[str, Any]:
    """Accept all tracked changes: <w:ins>→keep inner runs, <w:del>→remove."""
    from docx import Document
    from lxml import etree

    src = _resolve(file_id)
    t0 = time.time()
    doc = Document(str(src))
    body = doc.element.body
    n_ins = 0
    n_del = 0

    # Process <w:ins>: replace the element with its children (lift runs up).
    # list() materializes the iterator before we mutate the tree — without it,
    # nested <w:ins> (reply paragraphs produce these) get silently skipped.
    for ins in list(body.iter(qn("w:ins"))):
        parent = ins.getparent()
        idx = list(parent).index(ins)
        for child in list(ins):
            parent.insert(idx, child)
            idx += 1
        parent.remove(ins)
        n_ins += 1

    # Process <w:del>: remove entirely
    for d in list(body.iter(qn("w:del"))):
        parent = d.getparent()
        if parent is not None:
            parent.remove(d)
            n_del += 1

    fid, dst = _new_docx_output()
    doc.save(str(dst))
    return {
        "file_id": fid, "url": public_url(fid),
        "insertions_accepted": n_ins, "deletions_accepted": n_del,
        "ms_elapsed": int((time.time() - t0) * 1000),
    }


async def docx_reject_all_revisions(file_id: str) -> dict[str, Any]:
    """Reject all tracked changes: <w:ins>→remove, <w:del>→keep but rename
    delText back to t."""
    from docx import Document
    from lxml import etree

    src = _resolve(file_id)
    t0 = time.time()
    doc = Document(str(src))
    body = doc.element.body
    n_ins = 0
    n_del = 0

    # Reject <w:ins>: remove entirely
    for ins in list(body.iter(qn("w:ins"))):
        parent = ins.getparent()
        if parent is not None:
            parent.remove(ins)
            n_ins += 1

    # Reject <w:del>: lift inner runs, rename delText → t
    for d in list(body.iter(qn("w:del"))):
        parent = d.getparent()
        idx = list(parent).index(d)
        for child in list(d):
            for t in child.findall(qn("w:delText")):
                t.tag = qn("w:t")
            parent.insert(idx, child)
            idx += 1
        parent.remove(d)
        n_del += 1

    fid, dst = _new_docx_output()
    doc.save(str(dst))
    return {
        "file_id": fid, "url": public_url(fid),
        "insertions_rejected": n_ins, "deletions_rejected": n_del,
        "ms_elapsed": int((time.time() - t0) * 1000),
    }


async def docx_list_revisions(file_id: str) -> dict[str, Any]:
    """Enumerate all tracked changes in the document."""
    from docx import Document

    src = _resolve(file_id)
    doc = Document(str(src))
    body = doc.element.body
    revisions = []
    for el in body.iter():
        if el.tag == qn("w:ins"):
            kind = "insertion"
        elif el.tag == qn("w:del"):
            kind = "deletion"
        else:
            continue
        text_parts = []
        for t in el.iter():
            if t.tag in (qn("w:t"), qn("w:delText")):
                text_parts.append(t.text or "")
        revisions.append({
            "id": el.get(qn("w:id")),
            "kind": kind,
            "author": el.get(qn("w:author")),
            "date": el.get(qn("w:date")),
            "text": "".join(text_parts)[:300],
        })
    return {"file_id": file_id, "count": len(revisions), "revisions": revisions}


# ============================================================
# Comments
# ============================================================


_COMMENTS_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)
_COMMENTS_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)
_COMMENTS_EXT_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtended+xml"
)
_COMMENTS_EXT_REL_TYPE = (
    "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"
)
NS_W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
NS_W15 = "http://schemas.microsoft.com/office/word/2012/wordml"


def _qn_w14(tag: str) -> str:
    return f"{{{NS_W14}}}{tag}"


def _qn_w15(tag: str) -> str:
    return f"{{{NS_W15}}}{tag}"


def _ensure_para_id(p_elem) -> str:
    """Ensure a w:p element has a w14:paraId attribute; return its value."""
    import os
    pid = p_elem.get(_qn_w14("paraId"))
    if pid:
        return pid
    pid = os.urandom(4).hex().upper()
    p_elem.set(_qn_w14("paraId"), pid)
    return pid


def _ensure_comments_part(zip_path: Path) -> tuple[bytes, str]:
    """Read existing word/comments.xml or return a fresh empty one.
    Returns (xml_bytes, rid_to_use_or_create)."""
    pass  # complex zip-rewrite; we use a higher-level approach below


async def docx_add_comment(
    file_id: str, anchor: str, comment_text: str,
    *, author: str = "Kimi K2.6", initials: str = "K",
) -> dict[str, Any]:
    """Anchor a comment to the paragraph whose text equals `anchor`. Manages
    word/comments.xml, [Content_Types].xml registration, and document
    relationships if this is the first comment."""
    from lxml import etree
    import shutil

    src = _resolve(file_id)
    t0 = time.time()
    out_id = new_file_id("docx")
    out_path = output_path(out_id)
    shutil.copyfile(str(src), str(out_path))

    target = anchor.strip()
    date = _now_iso()

    # Open as zip, manipulate parts
    with zipfile.ZipFile(str(out_path), "r") as zin:
        names = zin.namelist()
        document_xml = zin.read("word/document.xml")
        try:
            comments_xml = zin.read("word/comments.xml")
        except KeyError:
            comments_xml = None
        try:
            content_types_xml = zin.read("[Content_Types].xml")
        except KeyError:
            content_types_xml = None
        try:
            doc_rels = zin.read("word/_rels/document.xml.rels")
        except KeyError:
            doc_rels = None

    # Parse document.xml
    doc_root = etree.fromstring(document_xml)
    body = doc_root.find(qn("w:body"))
    if body is None:
        raise ValueError("document.xml has no body")

    # Find anchor paragraph
    target_p = None
    for p in body.iter(qn("w:p")):
        # paragraph text = concat of all w:t descendants
        text_parts = [t.text or "" for t in p.iter(qn("w:t"))]
        if "".join(text_parts).strip() == target:
            target_p = p
            break
    if target_p is None:
        raise ValueError(f"anchor paragraph not found: {anchor!r}")

    # Determine new comment ID
    if comments_xml:
        comments_root = etree.fromstring(comments_xml)
        existing_ids = [int(c.get(qn("w:id")) or 0)
                        for c in comments_root.findall(qn("w:comment"))]
        new_cid = (max(existing_ids) + 1) if existing_ids else 0
    else:
        comments_root = etree.Element(
            qn("w:comments"), nsmap={"w": NS_W},
        )
        new_cid = 0

    # Add comment element
    c = etree.SubElement(comments_root, qn("w:comment"))
    c.set(qn("w:id"), str(new_cid))
    c.set(qn("w:author"), author)
    c.set(qn("w:date"), date)
    c.set(qn("w:initials"), initials)
    cp = etree.SubElement(c, qn("w:p"))
    cr = etree.SubElement(cp, qn("w:r"))
    ct = etree.SubElement(cr, qn("w:t"))
    ct.text = comment_text
    ct.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    # Insert range markers around the target paragraph's runs
    range_start = etree.Element(qn("w:commentRangeStart"))
    range_start.set(qn("w:id"), str(new_cid))
    range_end = etree.Element(qn("w:commentRangeEnd"))
    range_end.set(qn("w:id"), str(new_cid))
    ref_run = etree.Element(qn("w:r"))
    ref_run_rpr = etree.SubElement(ref_run, qn("w:rPr"))
    ref_style = etree.SubElement(ref_run_rpr, qn("w:rStyle"))
    ref_style.set(qn("w:val"), "CommentReference")
    ref_marker = etree.SubElement(ref_run, qn("w:commentReference"))
    ref_marker.set(qn("w:id"), str(new_cid))

    # Insert range_start at the beginning of paragraph's runs, range_end at the end,
    # and the reference run after range_end
    first_run_idx = None
    for i, child in enumerate(target_p):
        if child.tag == qn("w:r"):
            first_run_idx = i
            break
    if first_run_idx is None:
        target_p.append(range_start)
        target_p.append(range_end)
        target_p.append(ref_run)
    else:
        target_p.insert(first_run_idx, range_start)
        target_p.append(range_end)
        target_p.append(ref_run)

    # Update [Content_Types].xml if needed
    if content_types_xml:
        ct_root = etree.fromstring(content_types_xml)
        has_comments_override = any(
            o.get("PartName") == "/word/comments.xml"
            for o in ct_root.findall(f"{{{NS_CT}}}Override")
        )
        if not has_comments_override:
            override = etree.SubElement(ct_root, f"{{{NS_CT}}}Override")
            override.set("PartName", "/word/comments.xml")
            override.set("ContentType", _COMMENTS_CONTENT_TYPE)
        new_content_types_xml = etree.tostring(ct_root,
                                                xml_declaration=True,
                                                encoding="UTF-8",
                                                standalone=True)
    else:
        new_content_types_xml = content_types_xml

    # Update word/_rels/document.xml.rels if comments rel doesn't exist
    if doc_rels:
        rels_root = etree.fromstring(doc_rels)
        has_comments_rel = any(
            r.get("Type") == _COMMENTS_REL_TYPE
            for r in rels_root.findall(f"{{{NS_PR}}}Relationship")
        )
        if not has_comments_rel:
            existing_ids = [r.get("Id", "") for r in rels_root]
            n = max([int(re.match(r"rId(\d+)", x).group(1))
                     for x in existing_ids if re.match(r"rId(\d+)", x)] or [0])
            new_rel = etree.SubElement(rels_root, f"{{{NS_PR}}}Relationship")
            new_rel.set("Id", f"rId{n+1}")
            new_rel.set("Type", _COMMENTS_REL_TYPE)
            new_rel.set("Target", "comments.xml")
        new_doc_rels = etree.tostring(rels_root,
                                       xml_declaration=True,
                                       encoding="UTF-8",
                                       standalone=True)
    else:
        new_doc_rels = doc_rels

    # Write all updated parts back into the zip
    new_doc_xml = etree.tostring(doc_root,
                                  xml_declaration=True,
                                  encoding="UTF-8",
                                  standalone=True)
    new_comments_xml = etree.tostring(comments_root,
                                       xml_declaration=True,
                                       encoding="UTF-8",
                                       standalone=True)

    # Re-zip: copy everything except the parts we're replacing, then add updated parts
    tmp_path = out_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(str(out_path), "r") as zin, \
         zipfile.ZipFile(str(tmp_path), "w", zipfile.ZIP_DEFLATED) as zout:
        replace_set = {"word/document.xml", "word/comments.xml"}
        if new_content_types_xml:
            replace_set.add("[Content_Types].xml")
        if new_doc_rels:
            replace_set.add("word/_rels/document.xml.rels")
        for item in zin.infolist():
            if item.filename in replace_set:
                continue
            zout.writestr(item, zin.read(item.filename))
        zout.writestr("word/document.xml", new_doc_xml)
        zout.writestr("word/comments.xml", new_comments_xml)
        if new_content_types_xml:
            zout.writestr("[Content_Types].xml", new_content_types_xml)
        if new_doc_rels:
            zout.writestr("word/_rels/document.xml.rels", new_doc_rels)

    tmp_path.replace(out_path)
    return {
        "file_id": out_id, "url": public_url(out_id),
        "comment_id": new_cid, "author": author,
        "ms_elapsed": int((time.time() - t0) * 1000),
    }


async def docx_list_comments(file_id: str) -> dict[str, Any]:
    """List all comments in the document."""
    from lxml import etree

    src = _resolve(file_id)
    try:
        with zipfile.ZipFile(str(src), "r") as z:
            try:
                comments_xml = z.read("word/comments.xml")
            except KeyError:
                return {"file_id": file_id, "count": 0, "comments": []}
    except (zipfile.BadZipFile, KeyError):
        return {"file_id": file_id, "count": 0, "comments": []}

    root = etree.fromstring(comments_xml)
    out = []
    for c in root.findall(qn("w:comment")):
        text_parts = [t.text or "" for t in c.iter(qn("w:t"))]
        out.append({
            "id": c.get(qn("w:id")),
            "author": c.get(qn("w:author")),
            "initials": c.get(qn("w:initials")),
            "date": c.get(qn("w:date")),
            "text": "".join(text_parts),
        })
    return {"file_id": file_id, "count": len(out), "comments": out}


# ============================================================
# Comment threading + resolution (commentsExtended.xml)
# ============================================================


def _read_part(zin: zipfile.ZipFile, name: str) -> Optional[bytes]:
    try:
        return zin.read(name)
    except KeyError:
        return None


def _ensure_commentex_part(comments_ext_xml: Optional[bytes]):
    """Return commentsExtended root, parsing existing or creating fresh."""
    from lxml import etree
    if comments_ext_xml:
        return etree.fromstring(comments_ext_xml)
    return etree.Element(_qn_w15("commentsEx"), nsmap={"w15": NS_W15})


def _find_or_create_commentex(ext_root, para_id: str, *, parent_para_id: Optional[str] = None):
    """Find <w15:commentEx w15:paraId=...> or create one. Returns the element."""
    from lxml import etree
    for el in ext_root.findall(_qn_w15("commentEx")):
        if el.get(_qn_w15("paraId")) == para_id:
            return el
    el = etree.SubElement(ext_root, _qn_w15("commentEx"))
    el.set(_qn_w15("paraId"), para_id)
    el.set(_qn_w15("done"), "0")
    if parent_para_id:
        el.set(_qn_w15("paraIdParent"), parent_para_id)
    return el


async def docx_reply_to_comment(
    file_id: str, parent_comment_id: int, reply_text: str,
    *, author: str = "Kimi K2.6", initials: str = "K",
) -> dict[str, Any]:
    """Add a reply to comment `parent_comment_id`. Manages commentsExtended.xml
    threading via w15:paraIdParent. The reply itself has no body anchor — Word
    displays it nested under the parent in the review pane."""
    from lxml import etree
    import shutil

    src = _resolve(file_id)
    t0 = time.time()
    out_id = new_file_id("docx")
    out_path = output_path(out_id)
    shutil.copyfile(str(src), str(out_path))
    date = _now_iso()

    with zipfile.ZipFile(str(out_path), "r") as zin:
        comments_xml = _read_part(zin, "word/comments.xml")
        comments_ext_xml = _read_part(zin, "word/commentsExtended.xml")
        content_types_xml = _read_part(zin, "[Content_Types].xml")
        doc_rels_xml = _read_part(zin, "word/_rels/document.xml.rels")

    if not comments_xml:
        raise ValueError("document has no comments — cannot reply")

    comments_root = etree.fromstring(comments_xml)
    ext_root = _ensure_commentex_part(comments_ext_xml)

    # Locate parent comment + its paragraph's paraId
    parent_c = None
    for c in comments_root.findall(qn("w:comment")):
        if c.get(qn("w:id")) == str(parent_comment_id):
            parent_c = c
            break
    if parent_c is None:
        raise ValueError(f"parent comment id {parent_comment_id} not found")

    parent_p = parent_c.find(qn("w:p"))
    if parent_p is None:
        raise ValueError("parent comment has no paragraph")
    parent_para_id = _ensure_para_id(parent_p)
    _find_or_create_commentex(ext_root, parent_para_id)  # ensure parent registered

    # New reply comment
    existing_ids = [int(c.get(qn("w:id")) or 0)
                    for c in comments_root.findall(qn("w:comment"))]
    new_cid = max(existing_ids) + 1 if existing_ids else 0

    c = etree.SubElement(comments_root, qn("w:comment"))
    c.set(qn("w:id"), str(new_cid))
    c.set(qn("w:author"), author)
    c.set(qn("w:date"), date)
    c.set(qn("w:initials"), initials)
    cp = etree.SubElement(c, qn("w:p"))
    new_para_id = _ensure_para_id(cp)
    cr = etree.SubElement(cp, qn("w:r"))
    ct = etree.SubElement(cr, qn("w:t"))
    ct.text = reply_text
    ct.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    # Register reply in commentsExtended with paraIdParent
    _find_or_create_commentex(ext_root, new_para_id, parent_para_id=parent_para_id)

    # Content types: add commentsExtended override if missing
    new_content_types_xml = None
    if content_types_xml:
        ct_root = etree.fromstring(content_types_xml)
        has_ex = any(
            o.get("PartName") == "/word/commentsExtended.xml"
            for o in ct_root.findall(f"{{{NS_CT}}}Override")
        )
        if not has_ex:
            ov = etree.SubElement(ct_root, f"{{{NS_CT}}}Override")
            ov.set("PartName", "/word/commentsExtended.xml")
            ov.set("ContentType", _COMMENTS_EXT_CONTENT_TYPE)
            new_content_types_xml = etree.tostring(
                ct_root, xml_declaration=True, encoding="UTF-8", standalone=True,
            )

    # Relationships: add commentsExtended rel if missing
    new_doc_rels_xml = None
    if doc_rels_xml:
        rels_root = etree.fromstring(doc_rels_xml)
        has_ex_rel = any(
            r.get("Type") == _COMMENTS_EXT_REL_TYPE
            for r in rels_root.findall(f"{{{NS_PR}}}Relationship")
        )
        if not has_ex_rel:
            existing = [r.get("Id", "") for r in rels_root]
            n = max([int(re.match(r"rId(\d+)", x).group(1))
                     for x in existing if re.match(r"rId(\d+)", x)] or [0])
            rel = etree.SubElement(rels_root, f"{{{NS_PR}}}Relationship")
            rel.set("Id", f"rId{n+1}")
            rel.set("Type", _COMMENTS_EXT_REL_TYPE)
            rel.set("Target", "commentsExtended.xml")
            new_doc_rels_xml = etree.tostring(
                rels_root, xml_declaration=True, encoding="UTF-8", standalone=True,
            )

    new_comments_xml = etree.tostring(
        comments_root, xml_declaration=True, encoding="UTF-8", standalone=True,
    )
    new_ext_xml = etree.tostring(
        ext_root, xml_declaration=True, encoding="UTF-8", standalone=True,
    )

    tmp_path = out_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(str(out_path), "r") as zin, \
         zipfile.ZipFile(str(tmp_path), "w", zipfile.ZIP_DEFLATED) as zout:
        replace = {"word/comments.xml", "word/commentsExtended.xml"}
        if new_content_types_xml:
            replace.add("[Content_Types].xml")
        if new_doc_rels_xml:
            replace.add("word/_rels/document.xml.rels")
        for item in zin.infolist():
            if item.filename in replace:
                continue
            zout.writestr(item, zin.read(item.filename))
        zout.writestr("word/comments.xml", new_comments_xml)
        zout.writestr("word/commentsExtended.xml", new_ext_xml)
        if new_content_types_xml:
            zout.writestr("[Content_Types].xml", new_content_types_xml)
        if new_doc_rels_xml:
            zout.writestr("word/_rels/document.xml.rels", new_doc_rels_xml)
    tmp_path.replace(out_path)

    return {
        "file_id": out_id, "url": public_url(out_id),
        "reply_id": new_cid, "parent_id": parent_comment_id,
        "author": author,
        "ms_elapsed": int((time.time() - t0) * 1000),
    }


async def docx_resolve_comment(
    file_id: str, comment_id: int, *, resolved: bool = True,
) -> dict[str, Any]:
    """Mark a comment thread as resolved (w15:done='1') or reopened ('0')."""
    from lxml import etree
    import shutil

    src = _resolve(file_id)
    t0 = time.time()
    out_id = new_file_id("docx")
    out_path = output_path(out_id)
    shutil.copyfile(str(src), str(out_path))

    with zipfile.ZipFile(str(out_path), "r") as zin:
        comments_xml = _read_part(zin, "word/comments.xml")
        comments_ext_xml = _read_part(zin, "word/commentsExtended.xml")
        content_types_xml = _read_part(zin, "[Content_Types].xml")
        doc_rels_xml = _read_part(zin, "word/_rels/document.xml.rels")

    if not comments_xml:
        raise ValueError("document has no comments")

    comments_root = etree.fromstring(comments_xml)
    target_c = None
    for c in comments_root.findall(qn("w:comment")):
        if c.get(qn("w:id")) == str(comment_id):
            target_c = c
            break
    if target_c is None:
        raise ValueError(f"comment id {comment_id} not found")

    target_p = target_c.find(qn("w:p"))
    if target_p is None:
        raise ValueError("comment has no paragraph")
    para_id = _ensure_para_id(target_p)

    ext_root = _ensure_commentex_part(comments_ext_xml)
    cex = _find_or_create_commentex(ext_root, para_id)
    cex.set(_qn_w15("done"), "1" if resolved else "0")

    # Content_Types + rels updates if commentsExtended is new
    new_content_types_xml = None
    if content_types_xml and not comments_ext_xml:
        ct_root = etree.fromstring(content_types_xml)
        has_ex = any(
            o.get("PartName") == "/word/commentsExtended.xml"
            for o in ct_root.findall(f"{{{NS_CT}}}Override")
        )
        if not has_ex:
            ov = etree.SubElement(ct_root, f"{{{NS_CT}}}Override")
            ov.set("PartName", "/word/commentsExtended.xml")
            ov.set("ContentType", _COMMENTS_EXT_CONTENT_TYPE)
            new_content_types_xml = etree.tostring(
                ct_root, xml_declaration=True, encoding="UTF-8", standalone=True,
            )

    new_doc_rels_xml = None
    if doc_rels_xml and not comments_ext_xml:
        rels_root = etree.fromstring(doc_rels_xml)
        has_ex_rel = any(
            r.get("Type") == _COMMENTS_EXT_REL_TYPE
            for r in rels_root.findall(f"{{{NS_PR}}}Relationship")
        )
        if not has_ex_rel:
            existing = [r.get("Id", "") for r in rels_root]
            n = max([int(re.match(r"rId(\d+)", x).group(1))
                     for x in existing if re.match(r"rId(\d+)", x)] or [0])
            rel = etree.SubElement(rels_root, f"{{{NS_PR}}}Relationship")
            rel.set("Id", f"rId{n+1}")
            rel.set("Type", _COMMENTS_EXT_REL_TYPE)
            rel.set("Target", "commentsExtended.xml")
            new_doc_rels_xml = etree.tostring(
                rels_root, xml_declaration=True, encoding="UTF-8", standalone=True,
            )

    new_comments_xml = etree.tostring(
        comments_root, xml_declaration=True, encoding="UTF-8", standalone=True,
    )
    new_ext_xml = etree.tostring(
        ext_root, xml_declaration=True, encoding="UTF-8", standalone=True,
    )

    tmp_path = out_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(str(out_path), "r") as zin, \
         zipfile.ZipFile(str(tmp_path), "w", zipfile.ZIP_DEFLATED) as zout:
        replace = {"word/comments.xml", "word/commentsExtended.xml"}
        if new_content_types_xml:
            replace.add("[Content_Types].xml")
        if new_doc_rels_xml:
            replace.add("word/_rels/document.xml.rels")
        for item in zin.infolist():
            if item.filename in replace:
                continue
            zout.writestr(item, zin.read(item.filename))
        zout.writestr("word/comments.xml", new_comments_xml)
        zout.writestr("word/commentsExtended.xml", new_ext_xml)
        if new_content_types_xml:
            zout.writestr("[Content_Types].xml", new_content_types_xml)
        if new_doc_rels_xml:
            zout.writestr("word/_rels/document.xml.rels", new_doc_rels_xml)
    tmp_path.replace(out_path)

    return {
        "file_id": out_id, "url": public_url(out_id),
        "comment_id": comment_id, "resolved": resolved,
        "ms_elapsed": int((time.time() - t0) * 1000),
    }
